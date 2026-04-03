from docx import Document

doc = Document('VII_VIII_Report_Cards.docx')
print(f"Total tables: {len(doc.tables)}")

if len(doc.tables) > 0:
    table = doc.tables[0]
    print(f"Table 0 has {len(table.rows)} rows")
    for r_idx, row in enumerate(table.rows):
        cells_text = [c.text.strip()[:20].replace('\n', ' ') for c in row.cells]
        print(f"Row {r_idx}: {cells_text}")
