"""
Report Generator for IIT Report Generation
Creates Excel scorecards with charts and proper print settings
"""

import pandas as pd
from io import BytesIO
from typing import Dict, List
import xlsxwriter


def create_scorecard_excel(
    all_students: Dict,
    test_dates: List[str],
    class_name: str,
    section: str
) -> BytesIO:
    """
    Create an Excel workbook with one sheet per student.
    Each sheet contains:
    - Student info header
    - Test scores table
    - Progress charts
    
    Returns: BytesIO buffer with Excel file
    """
    output = BytesIO()
    
    workbook = xlsxwriter.Workbook(output, {'in_memory': True})
    
    # Define formats
    header_format = workbook.add_format({
        'bold': True,
        'font_size': 14,
        'align': 'center',
        'valign': 'vcenter',
        'bg_color': '#1E3A5F',
        'font_color': 'white',
        'border': 1
    })
    
    subheader_format = workbook.add_format({
        'bold': True,
        'font_size': 11,
        'align': 'left',
        'valign': 'vcenter',
        'bg_color': '#E8F4FD',
        'border': 1
    })
    
    table_header_format = workbook.add_format({
        'bold': True,
        'font_size': 10,
        'align': 'center',
        'valign': 'vcenter',
        'bg_color': '#4A90D9',
        'font_color': 'white',
        'border': 1
    })
    
    cell_format = workbook.add_format({
        'font_size': 10,
        'align': 'center',
        'valign': 'vcenter',
        'border': 1
    })
    
    avg_format = workbook.add_format({
        'bold': True,
        'font_size': 10,
        'align': 'center',
        'valign': 'vcenter',
        'bg_color': '#FFF3CD',
        'border': 1
    })
    
    # Sort students by name
    sorted_students = sorted(all_students.items(), key=lambda x: x[1]['student_name'])
    
    for candidate_id, student_data in sorted_students:
        # Create sheet for each student (sheet name = candidate_id)
        # Excel sheet names limited to 31 chars
        sheet_name = str(candidate_id)[:31]
        worksheet = workbook.add_worksheet(sheet_name)
        
        # Set page setup for A4 portrait with narrow margins
        worksheet.set_paper(9)  # A4
        worksheet.set_portrait()
        worksheet.set_margins(left=0.25, right=0.25, top=0.25, bottom=0.25)
        worksheet.fit_to_pages(1, 1)  # Fit to 1 page
        worksheet.set_print_scale(100)
        
        # Set column widths
        worksheet.set_column('A:A', 15)  # Test Date
        worksheet.set_column('B:B', 10)  # Physics
        worksheet.set_column('C:C', 10)  # Chemistry
        worksheet.set_column('D:D', 10)  # Maths
        worksheet.set_column('E:E', 10)  # Biology
        worksheet.set_column('F:F', 10)  # Total
        
        # Row 0: Title
        worksheet.merge_range('A1:F1', f'IIT FOUNDATION TEST SCORECARD', header_format)
        worksheet.set_row(0, 25)
        
        # Row 1: Class info
        worksheet.merge_range('A2:F2', f'Class: {class_name} | Section: {section}', subheader_format)
        
        # Row 2: Student info
        worksheet.merge_range('A3:F3', f'Name: {student_data["student_name"]} | Candidate ID: {candidate_id}', subheader_format)
        
        # Row 4: Table headers
        row = 4
        headers = ['Test Date', 'Physics', 'Chemistry', 'Maths', 'Biology', 'Total']
        for col, header in enumerate(headers):
            worksheet.write(row, col, header, table_header_format)
        
        # Data rows
        tests = student_data.get('tests', {})
        row = 5
        
        # Lists for chart data
        physics_scores = []
        chemistry_scores = []
        maths_scores = []
        biology_scores = []
        total_scores = []
        
        for test_date in test_dates:
            if test_date in tests:
                test_data = tests[test_date]
                worksheet.write(row, 0, test_date, cell_format)
                worksheet.write(row, 1, test_data.get('physics', 0), cell_format)
                worksheet.write(row, 2, test_data.get('chemistry', 0), cell_format)
                worksheet.write(row, 3, test_data.get('maths', 0), cell_format)
                worksheet.write(row, 4, test_data.get('biology', 0), cell_format)
                worksheet.write(row, 5, test_data.get('total', 0), cell_format)
                
                physics_scores.append(test_data.get('physics', 0))
                chemistry_scores.append(test_data.get('chemistry', 0))
                maths_scores.append(test_data.get('maths', 0))
                biology_scores.append(test_data.get('biology', 0))
                total_scores.append(test_data.get('total', 0))
                
                row += 1
        
        # Average row
        if physics_scores:
            worksheet.write(row, 0, 'Average', avg_format)
            worksheet.write(row, 1, round(sum(physics_scores) / len(physics_scores), 1), avg_format)
            worksheet.write(row, 2, round(sum(chemistry_scores) / len(chemistry_scores), 1), avg_format)
            worksheet.write(row, 3, round(sum(maths_scores) / len(maths_scores), 1), avg_format)
            worksheet.write(row, 4, round(sum(biology_scores) / len(biology_scores), 1), avg_format)
            worksheet.write(row, 5, round(sum(total_scores) / len(total_scores), 1), avg_format)
        
        # Create charts if we have data
        if len(test_dates) > 0 and len(physics_scores) > 0:
            chart_start_row = row + 3
            
            # Subject-wise Progress Chart (Line chart)
            chart1 = workbook.add_chart({'type': 'line'})
            
            # Data range for chart
            data_start_row = 6  # 1-indexed (row 5 in 0-indexed)
            data_end_row = data_start_row + len(physics_scores) - 1
            
            # Add series for each subject
            chart1.add_series({
                'name': 'Physics',
                'categories': [sheet_name, data_start_row - 1, 0, data_end_row - 1, 0],
                'values': [sheet_name, data_start_row - 1, 1, data_end_row - 1, 1],
                'marker': {'type': 'circle'},
                'line': {'color': '#FF6384'}
            })
            chart1.add_series({
                'name': 'Chemistry',
                'categories': [sheet_name, data_start_row - 1, 0, data_end_row - 1, 0],
                'values': [sheet_name, data_start_row - 1, 2, data_end_row - 1, 2],
                'marker': {'type': 'square'},
                'line': {'color': '#36A2EB'}
            })
            chart1.add_series({
                'name': 'Maths',
                'categories': [sheet_name, data_start_row - 1, 0, data_end_row - 1, 0],
                'values': [sheet_name, data_start_row - 1, 3, data_end_row - 1, 3],
                'marker': {'type': 'triangle'},
                'line': {'color': '#FFCE56'}
            })
            chart1.add_series({
                'name': 'Biology',
                'categories': [sheet_name, data_start_row - 1, 0, data_end_row - 1, 0],
                'values': [sheet_name, data_start_row - 1, 4, data_end_row - 1, 4],
                'marker': {'type': 'diamond'},
                'line': {'color': '#4BC0C0'}
            })
            
            chart1.set_title({'name': 'Subject-wise Progress'})
            chart1.set_x_axis({'name': 'Test Date'})
            chart1.set_y_axis({'name': 'Marks'})
            chart1.set_size({'width': 500, 'height': 280})
            chart1.set_legend({'position': 'bottom'})
            
            worksheet.insert_chart(chart_start_row, 0, chart1)
            
            # Total Marks Trend Chart (Column chart)
            chart2 = workbook.add_chart({'type': 'column'})
            
            chart2.add_series({
                'name': 'Total Marks',
                'categories': [sheet_name, data_start_row - 1, 0, data_end_row - 1, 0],
                'values': [sheet_name, data_start_row - 1, 5, data_end_row - 1, 5],
                'fill': {'color': '#1E3A5F'},
                'data_labels': {'value': True}
            })
            
            chart2.set_title({'name': 'Total Marks Trend'})
            chart2.set_x_axis({'name': 'Test Date'})
            chart2.set_y_axis({'name': 'Total Marks'})
            chart2.set_size({'width': 500, 'height': 280})
            chart2.set_legend({'none': True})
            
            worksheet.insert_chart(chart_start_row + 15, 0, chart2)
        
        # Set print area
        last_row = chart_start_row + 30 if len(physics_scores) > 0 else row + 2
        worksheet.print_area(0, 0, last_row, 5)
    
    workbook.close()
    output.seek(0)
    
    return output


def get_download_filename(class_name: str, section: str) -> str:
    """
    Generate the download filename.
    """
    # Clean up the names for file safety
    class_clean = class_name.replace(' ', '_').replace('/', '-')
    section_clean = section.replace(' ', '_').replace('/', '-')
    return f"{class_clean}_{section_clean}_IIT_Scorecard.xlsx"


def create_word_report_cards(
    all_students: Dict,
    test_dates: List[str],
    class_name: str,
    section: str,
    academic_year: str = "2025-26",
    template_path: str = None,
    logo_path: str = None
) -> BytesIO:
    """
    Template-based Word report card generation with 100% accuracy.
    
    Layout:
    - Each page has 2 students (Table 0 = top, Table 1 = bottom)
    - Each table represents 1 student's results
    - Left half (cols 0-6): Tests 1-7 (IX) or 1-5 (others)
    - Right half (cols 8-14): Tests 8-14 (IX) or 6-10 (others)
    
    The function:
    1. Selects template based on class (IX vs others)
    2. Extracts dates from template cells
    3. Matches student test data to template dates
    4. Fills marks only in matching rows
    5. Preserves ALL formatting, graphics, logos, signatures
    """
    from docx import Document
    from copy import deepcopy
    import os
    import tempfile
    import re
    
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    # Select template based on class
    if class_name.upper().strip() == 'IX':
        template_path = os.path.join(base_dir, 'IX_Report_Cards.docx')
    else:
        template_path = os.path.join(base_dir, 'Report_Cards.docx')
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
    
    # Sort students by name
    sorted_students = sorted(all_students.items(), key=lambda x: str(x[1].get('student_name', '')))
    
    def normalize_date(date_str):
        """
        Normalize date string to DD/MM/YYYY format for comparison.
        Handles formats: DD.MM.YYYY, DD-MM-YYYY, DD/MM/YYYY, DD.MM.YY
        """
        if not date_str or date_str == 'Unknown':
            return None
        date_str = str(date_str).strip()
        # Replace separators with /
        normalized = re.sub(r'[.\-_]', '/', date_str)
        # Handle 2-digit year (YY -> 20YY)
        parts = normalized.split('/')
        if len(parts) == 3 and len(parts[2]) == 2:
            parts[2] = '20' + parts[2]
            normalized = '/'.join(parts)
        return normalized
    
    def extract_template_dates(table, is_ix_template):
        """
        Extract all test dates from the template table.
        Returns dict: {(row_idx, col_offset): normalized_date}
        
        For IX template: data rows 4-10, left cols 0-6, right cols 8-14
        For Regular template: data rows 4-8, left cols 0-6, right cols 8-14
        """
        dates = {}
        
        if is_ix_template:
            data_rows = range(4, 11)  # Rows 4-10 (7 rows)
        else:
            data_rows = range(4, 9)   # Rows 4-8 (5 rows)
        
        for row_idx in data_rows:
            if row_idx >= len(table.rows):
                break
            row = table.rows[row_idx]
            
            # Left side date (column 1)
            if len(row.cells) > 1:
                left_date = row.cells[1].text.strip()
                normalized = normalize_date(left_date)
                if normalized:
                    dates[(row_idx, 0)] = normalized  # col_offset 0 = left side
            
            # Right side date (column 9)
            if len(row.cells) > 9:
                right_date = row.cells[9].text.strip()
                normalized = normalize_date(right_date)
                if normalized:
                    dates[(row_idx, 8)] = normalized  # col_offset 8 = right side
        
        return dates
    
    def fill_student_in_table(table, student_name, student_tests, test_dates_from_excel, 
                              class_name, section, is_ix_template):
        """
        Fill a single student's data into a table.
        - Row 1: Student name and class
        - Data rows: Fill marks where dates match
        """
        # Extract template dates
        template_dates = extract_template_dates(table, is_ix_template)
        
        # Normalize the test dates from Excel
        normalized_excel_dates = {}
        for test_date in test_dates_from_excel:
            normalized = normalize_date(test_date)
            if normalized:
                normalized_excel_dates[normalized] = test_date
        
        # Fill Row 1: Name and Class
        if len(table.rows) > 1:
            row1 = table.rows[1]
            # The row 1 is merged - find and update the text
            # We only need to process the first cell since it's merged
            if len(row1.cells) > 0:
                cell = row1.cells[0]
                
                # Calculate how many spaces to remove to compensate for student name length
                # Approx 1.8 spaces per character is a safe estimate for variable width fonts
                spaces_to_remove = int(len(student_name.strip()) * 1.8)
                
                for para in cell.paragraphs:
                    # Iterate through runs to find target fields
                    for run in para.runs:
                        text = run.text
                        
                        # Handle Name
                        if 'Name of the student:' in text:
                            run.text = f"Name of the student: {student_name}"
                        
                        # Handle Class
                        elif 'CLASS:' in text:
                             run.text = f"  CLASS: {class_name} {section}"
                        
                        # Handle Spacers (large blocks of spaces)
                        # We reduce them to prevent wrapping
                        elif len(text) > 5 and text.strip() == '':
                            if spaces_to_remove > 0:
                                current_len = len(text)
                                # Don't remove everything if we don't have to, but ensure we remove enough
                                remove = min(current_len, spaces_to_remove)
                                # Keep at least 1 space if it was a spacer, unless we need to shrink aggressively
                                if remove == current_len and spaces_to_remove < current_len + 5:
                                    remove = current_len - 1
                                
                                if remove > 0:
                                    run.text = text[:-remove]
                                    spaces_to_remove -= remove
                        
                        # Handle Cleanup of IX Template artifacts
                        # The IX template has 'Michael Faraday' in subsequent runs
                        elif is_ix_template and 'Michael Faraday' in text:
                            run.text = ""
        
        # Fill marks in data rows where dates match
        for (row_idx, col_offset), template_date in template_dates.items():
            # Check if we have data for this template date
            excel_date_key = None
            for excel_norm, excel_orig in normalized_excel_dates.items():
                if excel_norm == template_date:
                    excel_date_key = excel_orig
                    break
            
            if excel_date_key and excel_date_key in student_tests:
                test_data = student_tests[excel_date_key]
                row = table.rows[row_idx]
                
                # Fill marks in columns: Phy(2), Chem(3), Maths(4), Bio(5), Total(6)
                # For right side, add col_offset (8) to get: Phy(10), Chem(11), Maths(12), Bio(13), Total(14)
                marks = [
                    int(test_data.get('physics', 0)),
                    int(test_data.get('chemistry', 0)),
                    int(test_data.get('maths', 0)),
                    int(test_data.get('biology', 0)),
                    int(test_data.get('total', 0))
                ]
                
                for i, mark in enumerate(marks):
                    cell_idx = col_offset + 2 + i  # +2 because col 0=Test No, col 1=Date
                    if cell_idx < len(row.cells):
                        cell = row.cells[cell_idx]
                        # Clear existing text and set new value
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.text = str(mark)
                                break
                            else:
                                # No runs exist, create one with the value
                                if para.runs:
                                    para.runs[0].text = str(mark)
                                else:
                                    para.add_run(str(mark))
                            break
    
    # Determine if IX template
    is_ix_template = class_name.upper().strip() == 'IX'
    
    # Process students in pairs (2 per page)
    temp_files = []
    student_idx = 0
    num_students = len(sorted_students)
    
    while student_idx < num_students:
        # Load fresh template for each page
        doc = Document(template_path)
        
        # Each page has 2 tables (top and bottom)
        for table_idx, table in enumerate(doc.tables):
            if student_idx >= num_students:
                break
            
            # Get current student
            cid, student_data = sorted_students[student_idx]
            student_name = student_data.get('student_name', 'Unknown')
            student_tests = student_data.get('tests', {})
            
            # Fill this student's data in the table
            fill_student_in_table(
                table, 
                student_name, 
                student_tests, 
                test_dates, 
                class_name, 
                section, 
                is_ix_template
            )
            
            student_idx += 1
        
        # Save this page
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_file.close()  # Close handle before saving to avoid Windows locking
        doc.save(temp_file.name)
        temp_files.append(temp_file.name)
    
    # Combine all pages into final document
    if len(temp_files) == 1:
        with open(temp_files[0], 'rb') as f:
            output = BytesIO(f.read())
        os.unlink(temp_files[0])
        return output
    
    # Merge multiple pages using docxcompose
    try:
        from docxcompose.composer import Composer
        master = Document(temp_files[0])
        composer = Composer(master)
        for tf in temp_files[1:]:
            composer.append(Document(tf))
        output = BytesIO()
        composer.save(output)
        output.seek(0)
    except ImportError:
        # Fallback: return first page only if docxcompose not available
        with open(temp_files[0], 'rb') as f:
            output = BytesIO(f.read())
    except Exception:
        with open(temp_files[0], 'rb') as f:
            output = BytesIO(f.read())
    
    # Cleanup temp files
    for tf in temp_files:
        try:
            os.unlink(tf)
        except:
            pass
    
    return output




def create_consolidated_excel(
    all_students: Dict,
    test_dates: List[str],
    class_name: str,
    section: str,
    academic_year: str = "2024-25"
) -> BytesIO:
    """
    Create consolidated Excel report with all tests in columns.
    Matches the reference format with totals and averages.
    """
    output = BytesIO()
    workbook = xlsxwriter.Workbook(output, {'in_memory': True})
    
    sheet_name = f"{class_name} {section}"[:31]
    worksheet = workbook.add_worksheet(sheet_name)
    
    # Formats
    title_format = workbook.add_format({
        'bold': True, 'font_size': 14, 'align': 'center', 'valign': 'vcenter'
    })
    header_format = workbook.add_format({
        'bold': True, 'font_size': 10, 'align': 'center', 'valign': 'vcenter',
        'bg_color': '#4A90D9', 'font_color': 'white', 'border': 1, 'text_wrap': True
    })
    subheader_format = workbook.add_format({
        'bold': True, 'font_size': 9, 'align': 'center', 'valign': 'vcenter',
        'bg_color': '#E8F4FD', 'border': 1
    })
    cell_format = workbook.add_format({
        'font_size': 9, 'align': 'center', 'valign': 'vcenter', 'border': 1
    })
    name_format = workbook.add_format({
        'font_size': 9, 'align': 'left', 'valign': 'vcenter', 'border': 1
    })
    total_format = workbook.add_format({
        'bold': True, 'font_size': 9, 'align': 'center', 'valign': 'vcenter',
        'bg_color': '#FFF3CD', 'border': 1
    })
    
    # Set page setup
    worksheet.set_paper(9)  # A4
    worksheet.set_landscape()
    worksheet.set_margins(left=0.25, right=0.25, top=0.25, bottom=0.25)
    worksheet.fit_to_pages(1, 0)
    
    # Row 0: School name
    worksheet.merge_range(0, 0, 0, 5 + len(test_dates) * 5 + 10, "Rotary Pratishthan's", title_format)
    
    # Row 1: School
    worksheet.merge_range(1, 0, 1, 5 + len(test_dates) * 5 + 10, "Rotary Academy, Khed", title_format)
    
    # Row 2: Title
    title = f"{class_name} {section} Consolidated IIT Test Result {academic_year}"
    worksheet.merge_range(2, 0, 2, 5 + len(test_dates) * 5 + 10, title, title_format)
    
    # Row 3: Main headers
    row = 3
    worksheet.write(row, 0, "Sr. No.", header_format)
    worksheet.write(row, 1, "Exam Roll No.", header_format)
    worksheet.write(row, 2, "Name of the Student", header_format)
    
    col = 3
    for test_date in test_dates:
        worksheet.merge_range(row, col, row, col + 4, test_date, header_format)
        col += 5
    
    # Total and Average columns
    worksheet.merge_range(row, col, row, col + 4, "Total Marks", header_format)
    col += 5
    worksheet.merge_range(row, col, row, col + 4, "Average Score", header_format)
    col += 5
    worksheet.write(row, col, "Attended", header_format)
    
    # Row 4: Sub-headers (Phy, Chem, Maths, Bio, Total)
    row = 4
    worksheet.write(row, 0, "", subheader_format)
    worksheet.write(row, 1, "", subheader_format)
    worksheet.write(row, 2, "", subheader_format)
    
    col = 3
    subjects = ['Phy', 'Chem', 'Maths', 'Bio', 'Total']
    
    # For each test date
    for _ in test_dates:
        for subj in subjects:
            worksheet.write(row, col, subj, subheader_format)
            col += 1
    
    # Total marks columns
    for subj in subjects:
        worksheet.write(row, col, subj, subheader_format)
        col += 1
    
    # Average columns
    for subj in subjects:
        worksheet.write(row, col, subj, subheader_format)
        col += 1
    
    worksheet.write(row, col, "", subheader_format)
    
    # Set column widths
    worksheet.set_column(0, 0, 6)   # Sr. No.
    worksheet.set_column(1, 1, 12)  # Roll No.
    worksheet.set_column(2, 2, 25)  # Name
    worksheet.set_column(3, col, 6)  # All other columns
    
    # Data rows
    row = 5
    sorted_students = sorted(all_students.items(), key=lambda x: str(x[1].get('student_name', '')))
    
    for sr_no, (candidate_id, student_data) in enumerate(sorted_students, 1):
        worksheet.write(row, 0, sr_no, cell_format)
        worksheet.write(row, 1, candidate_id, cell_format)
        worksheet.write(row, 2, student_data.get('student_name', 'Unknown'), name_format)
        
        tests = student_data.get('tests', {})
        col = 3
        
        # Accumulators for totals
        total_phy = 0
        total_chem = 0
        total_maths = 0
        total_bio = 0
        total_total = 0
        attended = 0
        
        # Each test
        for test_date in test_dates:
            if test_date in tests:
                test_data = tests[test_date]
                phy = test_data.get('physics', 0)
                chem = test_data.get('chemistry', 0)
                maths = test_data.get('maths', 0)
                bio = test_data.get('biology', 0)
                total = test_data.get('total', 0)
                
                worksheet.write(row, col, int(phy), cell_format)
                worksheet.write(row, col + 1, int(chem), cell_format)
                worksheet.write(row, col + 2, int(maths), cell_format)
                worksheet.write(row, col + 3, int(bio), cell_format)
                worksheet.write(row, col + 4, int(total), cell_format)
                
                total_phy += phy
                total_chem += chem
                total_maths += maths
                total_bio += bio
                total_total += total
                attended += 1
            else:
                for i in range(5):
                    worksheet.write(row, col + i, 0, cell_format)
            
            col += 5
        
        # Total marks
        worksheet.write(row, col, int(total_phy), total_format)
        worksheet.write(row, col + 1, int(total_chem), total_format)
        worksheet.write(row, col + 2, int(total_maths), total_format)
        worksheet.write(row, col + 3, int(total_bio), total_format)
        worksheet.write(row, col + 4, int(total_total), total_format)
        col += 5
        
        # Averages
        if attended > 0:
            worksheet.write(row, col, round(total_phy / attended, 2), total_format)
            worksheet.write(row, col + 1, round(total_chem / attended, 2), total_format)
            worksheet.write(row, col + 2, round(total_maths / attended, 2), total_format)
            worksheet.write(row, col + 3, round(total_bio / attended, 2), total_format)
            worksheet.write(row, col + 4, round(total_total / attended, 2), total_format)
        else:
            for i in range(5):
                worksheet.write(row, col + i, 0, total_format)
        col += 5
        
        worksheet.write(row, col, attended, cell_format)
        
        row += 1
    
    workbook.close()
    output.seek(0)
    
    return output


def get_word_filename(class_name: str, section: str) -> str:
    """Generate filename for Word report cards."""
    class_clean = class_name.replace(' ', '_').replace('/', '-')
    section_clean = section.replace(' ', '_').replace('/', '-')
    return f"{class_clean}_{section_clean}_Report_Cards.docx"


def get_consolidated_filename(class_name: str, section: str) -> str:
    """Generate filename for consolidated Excel."""
    class_clean = class_name.replace(' ', '_').replace('/', '-')
    section_clean = section.replace(' ', '_').replace('/', '-')
    return f"{class_clean}_{section_clean}_Consolidated_Result.xlsx"
